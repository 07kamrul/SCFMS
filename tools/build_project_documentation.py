from __future__ import annotations

from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "documentation"
OUT.mkdir(exist_ok=True)

NAVY = "17324D"
BLUE = "2E74B5"
ORANGE = "EF6C00"
INK = "24313C"
MUTED = "65727E"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
WHITE = "FFFFFF"
RED = "9B1C1C"
GREEN = "276749"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run(run, size=None, color=INK, bold=None, italic=None, font="Aptos") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run(run, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 8),
        ("Subtitle", 13, MUTED, 0, 12),
        ("Heading 1", 17, NAVY, 16, 7),
        ("Heading 2", 13.5, BLUE, 12, 5),
        ("Heading 3", 11.5, NAVY, 9, 4),
    ):
        style = styles[name]
        style.font.name = "Aptos Display" if name in {"Title", "Heading 1"} else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.base_style = styles["Normal"]
    callout.font.name = "Aptos"
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = rgb(NAVY)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)


def add_header_footer(doc: Document, label: str) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(label.upper())
    set_run(r, size=8.5, color=MUTED, bold=True)
    p.paragraph_format.space_after = Pt(0)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("SCFMS  |  ")
    set_run(r, size=9, color=MUTED)
    add_page_field(fp)


def cover(doc: Document, kind: str, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    r = p.add_run("SCFMS")
    set_run(r, size=11, color=ORANGE, bold=True)
    p = doc.add_paragraph(style="Title")
    p.add_run(title)
    p = doc.add_paragraph(style="Subtitle")
    p.add_run(subtitle)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(kind.upper())
    set_run(r, size=10, color=WHITE, bold=True)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), NAVY)
    p_pr.append(shd)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(155)
    for label, value in (
        ("Repository", "SCFMS - Construction Field Management System"),
        ("Prepared from", "Current source code, migrations, tests, and product requirements"),
        ("Documentation date", date.today().isoformat()),
        ("Audience", "Developers, designers, QA, product owners, and implementation partners"),
    ):
        rr = p.add_run(f"{label}: ")
        set_run(rr, size=10, color=MUTED, bold=True)
        rr = p.add_run(value + "\n")
        set_run(rr, size=10, color=INK)
    doc.add_page_break()


def callout(doc: Document, label: str, text: str, color=NAVY) -> None:
    p = doc.add_paragraph(style="Callout")
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE)
    p_pr.append(shd)
    r = p.add_run(label + ": ")
    set_run(r, size=10.5, color=color, bold=True)
    r = p.add_run(text)
    set_run(r, size=10.5, color=INK)


def bullets(doc: Document, items: list[str], numbered=False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.add_run(item)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], font_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    set_table_geometry(t, widths)
    set_repeat_table_header(t.rows[0])
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_shading(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run(r, size=font_size, color=NAVY, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(str(value))
            set_run(r, size=font_size, color=INK)
    set_table_geometry(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_contents(doc: Document, sections: list[str]) -> None:
    doc.add_heading("How to use this document", level=1)
    doc.add_paragraph(
        "This is a repository-specific implementation guide. It describes what is present in the current codebase, "
        "how the pieces fit together, and what must still be configured or verified before production use."
    )
    doc.add_heading("Contents", level=2)
    bullets(doc, sections, numbered=True)
    callout(doc, "Status convention", "Present in code does not automatically mean production-ready. Items that require real devices, live infrastructure, legal review, or operational validation are called out explicitly.")
    doc.add_page_break()


def backend_doc() -> Path:
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc, "Backend Technical Documentation")
    cover(doc, "Technical reference and operations guide", "Backend Documentation", "FastAPI, PostgreSQL/PostGIS, Redis, MinIO, security, APIs, data model, testing, and deployment")
    sections = [
        "1. Product and backend scope", "2. Architecture", "3. Local setup and configuration",
        "4. Authentication, RBAC, and tenant isolation", "5. Data model", "6. API reference",
        "7. Core workflows", "8. Storage, offline support, and background jobs", "9. Testing",
        "10. Deployment and operations", "11. Security and privacy checklist", "12. Known gaps and roadmap",
    ]
    add_contents(doc, sections)

    doc.add_heading("1. Product and backend scope", level=1)
    doc.add_paragraph(
        "SCFMS is a mobile-first construction field management platform for small-to-mid-sized companies operating multiple sites. "
        "The backend supports company-separated accounts, map-based projects, workforce assignments, consent-based location tracking, tasks, issues, daily progress, media, notifications, and role-specific dashboards. It is intentionally not an ERP."
    )
    doc.add_heading("What is in scope", level=2)
    bullets(doc, [
        "Five fixed roles: Company Owner, HR Admin, Project Engineer, Site Engineer, and Employee.",
        "Strict company-level data isolation with tenant-scoped repositories and 404 responses for inaccessible records.",
        "OpenStreetMap-compatible project geometry stored with PostgreSQL/PostGIS.",
        "Location consent, tracking-hour checks, geofence classification, and access logging.",
        "Task and issue workflows with comments and photos; daily progress with stage entries and photo timelines.",
        "S3-compatible direct uploads using presigned URLs; MinIO is provided for local development.",
        "In-app notifications and device-token registration; dashboard summaries by role.",
    ])
    doc.add_heading("Explicit exclusions", level=2)
    bullets(doc, [
        "Payroll, accounting, leave, automated attendance, inventory/materials, machinery or vehicle tracking.",
        "Full chat, document management, safety management, and visitor management.",
        "Google Maps. The product requirement fixes OpenStreetMap as the map source.",
        "Phase-two items such as PDF/Excel exports, AI summaries, delay prediction, and a digital site diary.",
    ])

    doc.add_heading("2. Architecture", level=1)
    doc.add_heading("System context", level=2)
    table(doc, ["Layer", "Technology", "Responsibility"], [
        ["Mobile client", "Flutter / Dart", "Role-aware field UI, maps, camera, GPS, secure tokens, offline queues"],
        ["HTTP API", "FastAPI", "Versioned REST API, validation, dependency-based authorization, OpenAPI docs"],
        ["Domain layer", "Services", "Workflow rules, tenant checks, geofence logic, orchestration"],
        ["Persistence", "SQLAlchemy 2.0", "Repository pattern and transaction-aware data access"],
        ["Database", "PostgreSQL 16 + PostGIS 3.4", "Relational records and spatial project boundaries"],
        ["Cache/infra", "Redis 7", "Provisioned for caching/coordination; production use should be verified"],
        ["Object storage", "MinIO / S3", "Direct photo upload via presigned URLs"],
        ["Migrations", "Alembic", "Ordered schema evolution through revisions 0001-0008"],
    ], [1550, 2300, 5510], 8.3)
    doc.add_heading("Backend package boundaries", level=2)
    table(doc, ["Path", "Purpose"], [
        ["app/api/v1", "Thin route handlers and HTTP-level dependencies"],
        ["app/services", "Business rules and use-case orchestration"],
        ["app/repositories", "Company-scoped database queries"],
        ["app/models", "SQLAlchemy entities, relationships, and shared enums"],
        ["app/schemas", "Pydantic v2 request/response contracts"],
        ["app/permissions", "Auditable static role-to-permission matrix"],
        ["app/core", "Configuration, JWT/Argon2 security, exceptions, and logging"],
        ["app/utils", "Storage and geospatial utilities"],
        ["app/jobs", "Retention cleanup entry points"],
        ["alembic", "Database migrations"],
        ["tests", "API, isolation, RBAC, and feature tests"],
    ], [2300, 7060])
    callout(doc, "Request path", "Client request -> FastAPI router -> authentication/permission dependency -> service -> tenant-scoped repository -> PostgreSQL or S3 helper -> standard response envelope.")

    doc.add_heading("3. Local setup and configuration", level=1)
    doc.add_heading("Docker Compose quick start", level=2)
    bullets(doc, [
        "Copy the repository-root .env.example to .env if host ports need overriding.",
        "Copy backend/.env.example to backend/.env and replace JWT_SECRET_KEY and credentials.",
        "Review DATABASE_URL. The example remote URL is not appropriate for the bundled Docker database unless changed or removed so the POSTGRES_* fallback is used.",
        "Run: docker compose up --build",
        "Seed demo data once: docker compose exec backend python seed.py",
        "Open Swagger UI at http://localhost:18000/docs and health at http://localhost:18000/api/v1/health.",
    ], numbered=True)
    doc.add_heading("Default host ports", level=2)
    table(doc, ["Service", "Container", "Host default", "Override"], [
        ["Backend", "8000", "18000", "SCFMS_BACKEND_PORT"],
        ["PostgreSQL", "5432", "15432", "SCFMS_DB_PORT"],
        ["Redis", "6379", "16379", "SCFMS_REDIS_PORT"],
        ["MinIO API", "9000", "19000", "SCFMS_MINIO_PORT"],
        ["MinIO console", "9001", "19001", "SCFMS_MINIO_CONSOLE_PORT"],
    ], [1700, 1500, 1800, 4360])
    doc.add_heading("Important environment variables", level=2)
    table(doc, ["Group", "Variables", "Operational note"], [
        ["Application", "APP_NAME, ENVIRONMENT, DEBUG, API_V1_PREFIX", "Prefix must match the mobile API_BASE_URL."],
        ["CORS", "BACKEND_CORS_ORIGINS", "Relevant to browser clients; native Flutter Dio is not browser-CORS constrained."],
        ["JWT", "JWT_SECRET_KEY, JWT_ALGORITHM, token expiry values", "Use a long random secret; rotate deliberately."],
        ["Database", "DATABASE_URL or POSTGRES_*", "Application uses synchronous psycopg3; asyncpg URLs are normalized."],
        ["Storage", "S3_ENDPOINT_URL, S3_PUBLIC_URL, keys, bucket, region", "Public URL must be reachable from the device uploading media."],
        ["Protection", "MAX_LOGIN_ATTEMPTS, LOGIN_LOCKOUT_MINUTES", "Defaults are five attempts and fifteen minutes."],
        ["Geofence", "DEFAULT_NEAR_DISTANCE_METERS", "Default 300 m; company settings may override behavior."],
    ], [1300, 3500, 4560], 8.2)

    doc.add_heading("4. Authentication, RBAC, and tenant isolation", level=1)
    doc.add_heading("Authentication lifecycle", level=2)
    bullets(doc, [
        "Passwords are hashed with Argon2.",
        "Login returns a short-lived JWT access token and a database-backed refresh token.",
        "Refresh tokens are stored as SHA-256 hashes, rotated on use, and designed to detect reuse.",
        "The API supports logout of one session and logout of all sessions.",
        "Repeated login failures trigger a configurable account lockout.",
        "Authenticated users can retrieve their profile and change their password.",
    ])
    doc.add_heading("Role summary", level=2)
    table(doc, ["Role", "Primary responsibility", "Typical access"], [
        ["Company Owner", "Company oversight", "All permissions, settings, all projects, team, tracking, approvals, reports"],
        ["HR Admin", "Workforce administration", "Users, roles, assignments, all projects, assigned tracking view, reports"],
        ["Project Engineer", "Project coordination", "Assigned projects, tasks/approval, issues, progress view, team tracking"],
        ["Site Engineer", "Daily site execution", "Assigned projects, tasks/issues, progress submission, team tracking"],
        ["Employee", "Assigned field work", "Own projects, task updates, issue creation, photos, self tracking/share"],
    ], [1700, 2650, 5010], 8.4)
    callout(doc, "Authorization rule", "Mobile permission checks only shape the user interface. The backend permission dependency and company-scoped data access remain the security boundary.")
    doc.add_heading("Tenant isolation rules", level=2)
    bullets(doc, [
        "Every business record is associated directly or indirectly with company_id.",
        "Repository methods scope queries to the authenticated user's company.",
        "A record from another company is returned as not found, reducing information disclosure.",
        "Company creation is the bootstrap exception: it creates the company and initial owner account.",
        "Any new repository query must carry tenant scope; raw unscoped lookups should be treated as a security defect.",
    ])

    doc.add_heading("5. Data model", level=1)
    doc.add_paragraph("The current migrations define the following main aggregates and audit/support records.")
    table(doc, ["Aggregate", "Key records", "Purpose"], [
        ["Company", "Company, CompanySettings", "Tenant identity, near-distance and tracking policy/settings"],
        ["Identity", "User, RefreshToken, DeviceToken", "Roles, status, credentials, sessions, push-device registration"],
        ["Projects", "Project", "Name, status, progress and site polygon/location data"],
        ["Workforce", "Assignment", "User-project relationship, assignment role, start/end history, transfers"],
        ["Tracking", "LocationPoint, ActivityLog", "GPS samples, computed status, mock/accuracy context, auditable access/actions"],
        ["Tasks", "Task, TaskComment, TaskPhoto", "Assigned work, workflow, discussion and media"],
        ["Issues", "Issue, IssueStatusHistory, IssueComment, IssuePhoto", "Field problems, lifecycle history, discussion and evidence"],
        ["Progress", "DailyProgressReport, ProgressReportStageEntry, ProgressPhoto", "Daily site reporting, stage values and media timeline"],
        ["Notifications", "Notification", "User-targeted events and read state"],
    ], [1650, 3300, 4410], 8.2)
    doc.add_heading("Key enumerations", level=2)
    table(doc, ["Type", "Values"], [
        ["Project status", "planned, running, on_hold, delayed, completed, archived"],
        ["Task status", "todo, in_progress, blocked, submitted, approved, rejected, completed, cancelled"],
        ["Task priority", "low, medium, high, urgent"],
        ["Issue status", "open, assigned, in_progress, waiting, resolved, closed, reopened"],
        ["Issue priority", "low, medium, high, critical"],
        ["Issue category", "work_delay, design_problem, worker_shortage, site_access_problem, client_change, weather, quality_problem, utility_problem, approval_problem, other"],
        ["Location status", "inside_assigned, near_assigned, outside_assigned, inside_other_accessible, inside_other_unauthorized, no_assigned_project, location_disabled, offline, outside_tracking_hours, unknown"],
    ], [2050, 7310], 8.2)

    doc.add_heading("6. API reference", level=1)
    doc.add_paragraph("Base path defaults to /api/v1. Successful payloads use the shared Envelope schema. OpenAPI is available at /api/v1/openapi.json and interactive docs at /docs.")
    endpoints = [
        ("Health", "GET", "/health; /health/ready", "Liveness and dependency readiness"),
        ("Auth", "POST", "/auth/login; /auth/refresh", "Sign in and rotate tokens"),
        ("Auth", "POST", "/auth/logout; /auth/logout-all", "Revoke one or all sessions"),
        ("Auth", "GET/POST", "/auth/me; /auth/change-password", "Profile and credential maintenance"),
        ("Companies", "POST", "/companies/register", "Create company and initial owner"),
        ("Companies", "GET/PATCH", "/companies/settings", "Read or update tenant settings"),
        ("Users", "POST/GET", "/users", "Create and list users"),
        ("Users", "GET/PATCH", "/users/{user_id}", "Read or update a user"),
        ("Users", "POST", "/users/{id}/activate; /deactivate; /reset-password", "Account lifecycle operations"),
        ("Projects", "POST/GET", "/projects", "Create and list permitted projects"),
        ("Projects", "GET", "/projects/map; /projects/{project_id}", "Map dataset and detail"),
        ("Projects", "PATCH/POST/DELETE", "/projects/{id}; /archive; delete", "Edit, archive, or delete"),
        ("Assignments", "POST/GET", "/assignments; /assignments/me", "Manage or view assignment history"),
        ("Assignments", "POST", "/assignments/{id}/end; /transfer", "End or transfer assignment"),
        ("Locations", "POST", "/locations/consent; /locations", "Record consent and submit a location sample"),
        ("Locations", "GET", "/locations/me; /locations/team", "Own status and permitted team map"),
        ("Tasks", "POST/GET/PATCH", "/tasks; /tasks/{task_id}", "Create, list, read, and update tasks"),
        ("Tasks", "GET/POST", "/tasks/{id}/comments; /photos", "Task collaboration and evidence"),
        ("Issues", "POST/GET/PATCH", "/issues; /issues/{issue_id}", "Create, list, read, and update issues"),
        ("Issues", "GET", "/issues/{id}/history", "Status change audit trail"),
        ("Issues", "GET/POST", "/issues/{id}/comments; /photos", "Issue collaboration and evidence"),
        ("Progress", "POST/GET", "/progress-reports", "Create and list daily reports"),
        ("Progress", "GET", "/progress-reports/timeline; /{report_id}", "Photo timeline and report detail"),
        ("Progress", "GET/POST", "/progress-reports/{id}/photos", "Report media"),
        ("Uploads", "POST", "/uploads/presign", "Create direct-to-object-storage upload contract"),
        ("Notifications", "GET/POST", "/notifications; /{id}/read", "List notifications and mark read"),
        ("Notifications", "POST", "/notifications/device-tokens", "Register Android/iOS push token"),
        ("Dashboard", "GET", "/dashboard/me", "Role-specific summary for current user"),
    ]
    table(doc, ["Area", "Method", "Path(s)", "Purpose"], [list(x) for x in endpoints], [1250, 1350, 3900, 2860], 7.7)
    callout(doc, "Canonical contract", "Use the generated OpenAPI document for exact request fields, validation limits, query parameters, response shapes, and status codes. This guide is the architectural index, not a substitute for the machine-readable contract.")

    doc.add_heading("7. Core workflows", level=1)
    workflows = [
        ("Company onboarding", ["Register company and initial owner", "Log in", "Configure settings", "Create users and assign roles", "Create projects and boundaries", "Assign workers and engineers"]),
        ("Field execution", ["Engineer creates and assigns task", "Employee starts or updates task", "Employee attaches photos/comments", "Employee submits", "Authorized engineer approves or rejects", "Dashboard and notifications reflect changes"]),
        ("Issue handling", ["User records category, priority, description, project and optional assignee", "Evidence photos/comments are added", "Status changes append history", "Issue is resolved, closed, or reopened"]),
        ("Location tracking", ["Worker grants explicit consent", "Client starts visible foreground/background tracking within policy", "GPS points are submitted or queued", "Server calculates geofence status", "Authorized managers read the team map; worker sees own status"]),
        ("Daily progress", ["Site Engineer selects project/date", "Records stage-wise progress and notes", "Submits report", "Uploads photos through presigned storage", "Managers review reports and project photo timeline"]),
    ]
    for name, steps in workflows:
        doc.add_heading(name, level=2)
        bullets(doc, steps, numbered=True)

    doc.add_heading("8. Storage, offline support, and background jobs", level=1)
    doc.add_heading("Photo upload sequence", level=2)
    bullets(doc, [
        "Client asks POST /uploads/presign for an upload target.",
        "Backend creates a tenant-aware object key and returns a temporary S3-compatible upload URL.",
        "Client compresses and uploads the file directly to object storage.",
        "Client associates the resulting object key with a task, issue, or progress report.",
        "Failed mobile uploads are retained in a local retry queue and retried when connectivity returns.",
    ], numbered=True)
    doc.add_heading("Retention", level=2)
    doc.add_paragraph("app/jobs/retention_cleanup.py provides the location-retention cleanup entry point. Production must schedule it explicitly (for example, a cron job, worker, or platform scheduler), observe failures, and align the retention window with company policy and applicable law.")
    doc.add_heading("Redis", level=2)
    doc.add_paragraph("Redis is provisioned by Compose and configured in Settings. Confirm the exact runtime responsibilities required for production (notification cooldown/dedup, caching, queues, or locks) before sizing or treating it as a hard availability dependency.")

    doc.add_heading("9. Testing", level=1)
    doc.add_paragraph("The test suite includes authentication, RBAC, company isolation, dashboards, projects, assignments, location, tasks, issues, progress reports, uploads, and notifications. Tests use a separate database name on the same PostgreSQL/PostGIS server.")
    bullets(doc, [
        "Create scfms_test on the reachable PostgreSQL server.",
        "Set TEST_DATABASE_URL to a psycopg URL for scfms_test.",
        "Run pytest from backend/.",
        "Never point TEST_DATABASE_URL at seeded development or production data; fixtures reset schema state.",
    ], numbered=True)
    doc.add_heading("Recommended CI gates", level=2)
    bullets(doc, [
        "pytest with PostgreSQL/PostGIS service", "Alembic upgrade from an empty database", "Static formatting/linting and type checks",
        "OpenAPI schema generation/diff review", "Dependency and container vulnerability scans", "Tenant-isolation regression tests for every new resource",
    ])

    doc.add_heading("10. Deployment and operations", level=1)
    doc.add_heading("Production readiness checklist", level=2)
    bullets(doc, [
        "Set ENVIRONMENT=production and DEBUG=false.", "Generate and securely store a strong JWT_SECRET_KEY.",
        "Use managed or backed-up PostgreSQL with PostGIS enabled; apply Alembic migrations before serving traffic.",
        "Use private database/Redis networking and TLS at the ingress/reverse proxy.", "Use durable S3-compatible storage and a device-reachable public upload endpoint.",
        "Replace staging/production placeholder hostnames in mobile environment files.", "Configure only required browser CORS origins; do not use wildcard production origins.",
        "Schedule retention cleanup and database/object-store backups; test restore procedures.", "Centralize logs and add health, latency, error-rate, storage, and database alerts.",
        "Review token rotation, account recovery, key rotation, and incident-response procedures.",
    ])
    doc.add_heading("Runbook essentials", level=2)
    table(doc, ["Symptom", "Check first", "Likely action"], [
        ["API unavailable", "Container status, /health, logs", "Restart only after identifying dependency/config failure"],
        ["Readiness fails", "Database connectivity/migrations", "Restore network/credentials or complete migration"],
        ["Uploads fail", "S3 endpoint, bucket, public URL, device reachability", "Correct storage policy/URL; preserve client retry queue"],
        ["Mobile cannot connect", "API_BASE_URL, scheme/TLS, host routing", "Use emulator alias locally; replace placeholder remote host"],
        ["Unexpected 404", "Tenant, assignment, permission, resource id", "Confirm authorization before assuming missing data"],
        ["Tracking stale", "Consent, tracking hours, device permission, queue, last seen", "Surface accurate state; avoid silently forcing tracking"],
    ], [1800, 3400, 4160], 8.2)

    doc.add_heading("11. Security and privacy checklist", level=1)
    bullets(doc, [
        "Keep authentication and authorization server-side; never trust role claims from the UI alone.",
        "Add company scope to every new query and test cross-company denial.",
        "Validate upload content type, size, object-key ownership, and association permissions.",
        "Rate-limit login, refresh, presign, and high-volume location endpoints at the edge or application layer.",
        "Retain only the location precision and duration needed; log manager access to location data.",
        "Show tracking consent, current status, working-hour behavior, and revocation controls clearly in the mobile UI.",
        "Resolve launch-country privacy obligations and obtain legal review before a pilot involving workers.",
        "Do not log passwords, bearer tokens, refresh tokens, presigned URLs, or precise coordinates unnecessarily.",
    ])

    doc.add_heading("12. Known gaps and roadmap", level=1)
    table(doc, ["Item", "Current evidence", "Next decision/action"], [
        ["README status", "Milestone list says only foundation is complete, but routes/models/tests exist for later modules.", "Update README after end-to-end verification."],
        ["Staging/production", "Mobile hostnames are TBD placeholders.", "Provision environments and verify real-device connectivity."],
        ["Push delivery", "Device-token API and notifications exist.", "Confirm provider credentials, delivery worker, cooldown, and receipt behavior."],
        ["Offline conflicts", "Location/photo submission queues exist on mobile.", "Define conflict rules for editable drafts and repeated submissions."],
        ["Privacy/legal", "Consent, tracking policy, audit and retention concepts exist.", "Set market-specific policy and complete legal review."],
        ["Pilot validation", "Product metrics remain assumptions in the PRD.", "Recruit pilot, baseline metrics, and instrument adoption/latency."],
        ["Phase 2", "Exports, site diary, announcements, AI are explicitly excluded from MVP.", "Prioritize only after pilot evidence."],
    ], [1800, 3900, 3660], 8.1)

    path = OUT / "SCFMS_Backend_Documentation.docx"
    doc.save(path)
    return path


def mobile_doc() -> Path:
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc, "Mobile UI and Prototype Documentation")
    cover(doc, "UX specification and prototype blueprint", "Mobile App & UI Prototype Documentation", "Flutter architecture, role journeys, screen specifications, design tokens, components, states, and a Figma-ready build plan")
    sections = [
        "1. Product experience", "2. Best place to prototype", "3. Users, roles, and navigation",
        "4. Information architecture", "5. Design system", "6. Screen-by-screen specification",
        "7. Critical prototype flows", "8. Responsive behavior", "9. Offline, loading, empty, and error states",
        "10. Prototype data kit", "11. Figma build procedure", "12. Flutter handoff and acceptance checklist",
    ]
    add_contents(doc, sections)

    doc.add_heading("1. Product experience", level=1)
    doc.add_paragraph(
        "The SCFMS mobile app is the field-facing surface for company setup, projects, team assignment, live site tracking, task execution, issue reporting, daily progress, photos, notifications, and role-specific oversight. The intended experience is practical and high-visibility: readable outdoors, fast to scan, tolerant of weak connectivity, and explicit about location privacy."
    )
    callout(doc, "Design direction", "Material 3 with a construction-site character: safety orange for primary actions and attention, blue-grey for structure, strong status semantics, 8/12/16 px radii, generous tap targets, and restrained motion.")
    doc.add_heading("Experience principles", level=2)
    bullets(doc, [
        "Answer 'what needs my attention today?' on the first authenticated screen.",
        "Keep the current project, assignee, status, priority, and last-seen time visible where decisions depend on them.",
        "Never hide or disguise location tracking; consent and active state must be obvious.",
        "Make the important action reachable with one thumb and one clear primary button.",
        "Design every submit action for interruption, weak connectivity, retry, and duplicate prevention.",
        "Use status text plus color/icon; color alone is not sufficient.",
    ])

    doc.add_heading("2. Best place to prototype", level=1)
    doc.add_heading("Recommended: Figma", level=2)
    doc.add_paragraph(
        "Use Figma for the first clickable prototype. It is the easiest fit for this project because the work is primarily screen design and role-based user-flow validation, not yet another implementation. Build reusable components and variables, connect frames into flows, share a link, and preview the prototype on a real phone."
    )
    table(doc, ["Option", "Best use", "Trade-off", "Recommendation"], [
        ["Figma", "Wireframes through high-fidelity clickable mobile flows", "Prototype logic is simulated; final Flutter implementation remains separate", "Primary choice"],
        ["Penpot", "Open-source/self-hosted design collaboration", "Smaller ecosystem and handoff workflow may need more setup", "Use if open-source ownership is required"],
        ["FlutterFlow", "Higher-function prototype or rapid Flutter-oriented proof", "Can diverge from the existing BLoC/get_it/go_router codebase and create parallel code", "Use only after UX validation"],
        ["Flutter code", "Production behavior, maps, GPS, offline queues", "Slower for early exploration and stakeholder feedback", "Use for validated screens"],
    ], [1500, 3000, 3000, 1860], 8.1)
    doc.add_paragraph("Official references: Figma prototyping - https://www.figma.com/prototyping/ ; Figma mobile preview - https://help.figma.com/hc/en-us/articles/360040321093-View-prototypes-on-a-mobile-device")
    doc.add_heading("Prototype fidelity stages", level=2)
    bullets(doc, [
        "Stage 1 - grayscale flow: navigation, hierarchy, required fields, and success/error paths.",
        "Stage 2 - component system: color, typography, spacing, controls, status chips, cards, and map overlays.",
        "Stage 3 - high-fidelity role flows: realistic data, overlays, transitions, offline states, and device preview.",
        "Stage 4 - Flutter proof: only the map/tracking/offline interactions that a click-through prototype cannot validate.",
    ], numbered=True)

    doc.add_heading("3. Users, roles, and navigation", level=1)
    table(doc, ["Role", "Home priority", "Visible areas"], [
        ["Company Owner", "Company health and exceptions", "Home, Projects, Team, Tasks, Issues, Tracking, Profile"],
        ["HR Admin", "Workforce and assignments", "Home, Projects, Team, Tracking, Profile"],
        ["Project Engineer", "Approvals, blockers, project progress", "Home, Projects, Team, Tasks, Issues, Tracking, Profile"],
        ["Site Engineer", "Today's field execution and daily report", "Home, Projects, Team, Tasks, Issues, Tracking, Profile"],
        ["Employee", "Assigned work and tracking status", "Home, Projects, Tasks, Issues, Tracking, Profile"],
    ], [1650, 3100, 4610], 8.3)
    doc.add_heading("Navigation behavior", level=2)
    bullets(doc, [
        "Phone: Material NavigationBar at the bottom, generated from the authenticated role's permissions.",
        "Tablet/expanded width: NavigationRail on the left with persistent labels.",
        "Feature details and forms open through in-feature navigation, then return to the originating list/detail.",
        "Unauthenticated routes are Splash, Login, and Register Company. Authenticated users redirect to Dashboard.",
        "Notifications are accessed from the dashboard/app bar rather than occupying a permanent primary tab.",
    ])

    doc.add_heading("4. Information architecture", level=1)
    table(doc, ["Primary area", "Destinations and secondary screens"], [
        ["Authentication", "Splash -> Login -> Register Company; Change Password from Profile"],
        ["Home", "Role-specific dashboard -> notification list -> relevant task/issue/project"],
        ["Projects", "All Projects or My Projects -> Project Map/Detail -> Create/Edit -> Progress Reports -> Photo Timeline"],
        ["Team", "Team List -> User Detail -> Create/Edit User -> Assignment Form"],
        ["Tasks", "Task List -> Task Detail -> Create/Edit -> Comments/Photos -> Submit/Approve/Reject"],
        ["Issues", "Issue List -> Issue Detail -> Create/Edit -> History/Comments/Photos"],
        ["Tracking", "Consent & Own Status -> Team Map when permitted"],
        ["Profile", "Profile -> Change Password -> Company Settings when permitted"],
    ], [1900, 7460], 8.5)

    doc.add_heading("5. Design system", level=1)
    doc.add_heading("Foundation tokens", level=2)
    table(doc, ["Token", "Current/source-aligned value", "Prototype use"], [
        ["Primary", "Safety orange #EF6C00", "Primary buttons, active navigation, key attention"],
        ["Secondary light", "Blue-grey #37474F", "Structure, secondary emphasis"],
        ["Secondary dark", "Blue-grey #90A4AE", "Dark-mode secondary emphasis"],
        ["Radius", "8 / 12 / 16 px", "Inputs/buttons; cards; large surfaces/dialogs"],
        ["Spacing", "4 px base rhythm; favor 8, 12, 16, 24, 32", "Component gaps and page padding"],
        ["Motion", "150 / 250 / 400 ms, ease-out cubic", "Feedback, content change, larger transitions"],
        ["Typography", "Material 3 sizes; 600-700 headings; body line height 1.4", "Readable hierarchy and outdoor scanning"],
        ["Phone breakpoint", "Use existing app_breakpoints.dart as source", "Bottom navigation and single-column forms"],
    ], [1900, 3260, 4200], 8.3)
    doc.add_heading("Semantic status system", level=2)
    table(doc, ["Meaning", "Suggested color", "Examples"], [
        ["Positive/complete", "Green", "Inside assigned site, approved, completed, resolved"],
        ["In progress/info", "Blue", "Running project, in progress, submitted"],
        ["Warning", "Amber", "Near site, waiting, medium/high attention"],
        ["Critical", "Red", "Wrong site, delayed, blocked, overdue, critical issue"],
        ["Inactive/unknown", "Grey", "Offline, archived, cancelled, unavailable"],
        ["Primary action", "Safety orange", "Create, save, start tracking, submit"],
    ], [1800, 2200, 5360], 8.4)
    doc.add_heading("Reusable components to create first", level=2)
    bullets(doc, [
        "App shell: phone bottom navigation and tablet navigation rail variants.",
        "Top app bar: title, back, notification badge, overflow actions.",
        "Buttons: filled, outlined, text, destructive; loading and disabled states.",
        "Fields: text, password, dropdown, date/time, numeric percentage, search, multiline notes.",
        "StatusBadge: project/task/issue/location variants with label, icon, and semantic color.",
        "Summary card and metric tile; empty-state panel; error panel; skeleton/loading view.",
        "List row: avatar/icon, title, two metadata lines, status, chevron, optional trailing action.",
        "Photo picker/upload tile: empty, local preview, compressing, queued, uploading, failed, complete.",
        "Map controls: current location, fit bounds, layers/legend, project polygon, worker marker, accuracy halo.",
        "Consent panel and persistent tracking indicator.",
    ])

    doc.add_heading("6. Screen-by-screen specification", level=1)
    screen_rows = [
        ["Splash", "Restore secure session and route", "Centered mark/name, compact progress indicator", "Loading; session restored; session expired"],
        ["Login", "Authenticate", "Email, password, primary Sign In, Register Company link", "Validation, lockout, offline, invalid credentials"],
        ["Register Company", "Bootstrap tenant + owner", "Company details, owner details, password rules, Create Company", "Validation, duplicate, success redirect"],
        ["Dashboard", "Role-specific attention view", "Greeting, tracking state, metrics, urgent tasks/issues, project progress, shortcuts", "Loading, empty role state, partial-card failure"],
        ["Notifications", "Review events", "Unread emphasis, time, type icon, tap target, mark-read behavior", "Empty, pagination/loading, failed refresh"],
        ["Projects / My Projects", "Find permitted projects", "Search/filter, status chip, progress, dates, map entry, create FAB if allowed", "Empty, archived filter, offline cache"],
        ["Project Map", "See or edit boundary", "OSM tiles, polygon, fit control, status legend; edit handles in form mode", "No polygon, tile failure, location denied"],
        ["Project Form", "Create/edit project", "Identity, dates, status, progress, boundary drawing and validation", "Unsaved changes, invalid polygon, submit failure"],
        ["Team List", "Find and manage people", "Search, role/status filters, avatar initials, project count, add action", "No matches, inactive users, permission-limited"],
        ["User Detail", "Review identity and assignment history", "Contact, role/status, assignments, edit/deactivate/reset actions", "Confirmation, protected action, error"],
        ["User Form", "Create/edit account", "Name, email, phone, role, temporary password/create action", "Duplicate email, permission error"],
        ["Assignment Form", "Assign/transfer worker", "User, project, assignment role, start date, transfer/end context", "Conflict, confirmation, success"],
        ["Task List", "Prioritize work", "My/all toggle where allowed, project/status/priority filters, overdue marker, create FAB", "Empty, stale/offline, retry"],
        ["Task Detail", "Execute and review task", "Header status, assignee/project/due date, description, comments, photos, workflow actions", "Action confirmation, queued photo, conflict"],
        ["Task Form", "Create/edit/reassign", "Title, project, assignee, priority, due date, description", "Required fields, invalid date, submit failure"],
        ["Issue List", "Surface blockers", "Priority/category/status filters, critical prominence, create FAB", "Empty, offline, filter no-results"],
        ["Issue Detail", "Resolve with audit trail", "Summary, project/assignee, history, comments, photos, status action", "Reopen, failed upload, permission-limited"],
        ["Issue Form", "Report field issue quickly", "Project, category, priority, title/description, assignee, initial photo", "Draft/queued, validation, submit success"],
        ["Progress Reports", "Review daily site reporting", "Date, author, overall/stage progress, notes, photo count", "No report today, filter, offline"],
        ["Progress Report Form", "Submit daily report", "Project/date, stage entries, progress values, notes, photos, submit", "Draft recovery, duplicate date, queued media"],
        ["Photo Timeline", "Browse project evidence", "Date grouping, image grid, captions/metadata, full-screen viewer", "Empty, placeholder, load-more"],
        ["Tracking Consent", "Control and understand tracking", "Consent copy, permission state, tracking-hours note, start/stop, visible indicator, own status", "Denied, outside hours, GPS off, queued/offline"],
        ["Team Map", "Monitor permitted staff", "OSM map, project polygons, clustered/status markers, legend, filter, last-seen sheet", "No permission, no locations, stale/offline"],
        ["Profile", "Account and app controls", "Identity, role/company, theme/settings, change password, sign out", "Offline-safe display, sign-out confirmation"],
        ["Company Settings", "Configure tenant policies", "Near distance, tracking days/hours, retention-related settings, save", "Validation, unsaved changes, permission denied"],
        ["Change Password", "Rotate credentials", "Current/new/confirm, requirements, success sign-in behavior", "Mismatch, weak password, expired session"],
    ]
    table(doc, ["Screen", "User goal", "Required content/actions", "States to prototype"], screen_rows, [1500, 2000, 3860, 2000], 7.2)

    doc.add_heading("Screen composition notes", level=2)
    doc.add_heading("Dashboard", level=3)
    bullets(doc, [
        "Top zone: greeting, role/company, notification action, network/tracking state when relevant.",
        "First card must be role-specific: Owner sees company exceptions; Site Engineer sees today's project/report; Employee sees current task and own tracking status.",
        "Use two-column metric tiles only where width permits; stack on compact phones.",
        "Every exception card deep-links to the filtered source list or detail.",
    ])
    doc.add_heading("Maps", level=3)
    bullets(doc, [
        "Keep controls away from system gestures and the bottom navigation safe area.",
        "Use status-colored polygon borders/fills with sufficient transparency so OSM labels remain readable.",
        "Worker marker selection opens a bottom sheet with name, project, status, accuracy, and last seen.",
        "Do not imply exact presence when accuracy is poor or data is stale; show confidence and timestamp.",
    ])
    doc.add_heading("Forms", level=3)
    bullets(doc, [
        "One clear page title and one sticky or bottom-safe primary action.",
        "Group fields into identity, assignment/context, timing/status, details, and evidence.",
        "Preserve draft input through navigation or connection loss; confirm before discarding.",
        "Keep validation beside the field and summarize only when it helps recovery.",
    ])

    doc.add_heading("7. Critical prototype flows", level=1)
    flows = [
        ("Flow A - Company onboarding", "Owner", ["Splash", "Login", "Register Company", "Dashboard", "Company Settings", "Projects List", "Project Form + polygon", "Team List", "User Form", "Assignment Form"]),
        ("Flow B - Employee field day", "Employee", ["Login", "Dashboard", "Tracking Consent", "Own tracking active", "Tasks List", "Task Detail", "Add photo/comment", "Submit task", "Issue Form", "Offline queued confirmation"]),
        ("Flow C - Site daily report", "Site Engineer", ["Dashboard", "My Projects", "Progress Reports", "Progress Report Form", "Add stages/photos", "Queued upload or success", "Photo Timeline"]),
        ("Flow D - Manager exception response", "Project Engineer", ["Dashboard alert", "Issue Detail", "Review history/photos", "Assign/update status", "Team Map", "Worker detail sheet", "Task Detail", "Approve/reject"]),
        ("Flow E - HR assignment transfer", "HR Admin", ["Team List", "User Detail", "Assignment Form", "End/transfer confirmation", "Updated assignment history", "Team Map"]),
    ]
    for title, role, frames in flows:
        doc.add_heading(title, level=2)
        doc.add_paragraph(f"Primary role: {role}")
        bullets(doc, frames, numbered=True)
    callout(doc, "Prototype rule", "Each flow needs at least one happy path, one validation/error state, and one offline or permission-related recovery path. Do not prototype only ideal taps.")

    doc.add_heading("8. Responsive behavior", level=1)
    table(doc, ["Area", "Compact phone", "Large phone / landscape", "Tablet / expanded"], [
        ["Navigation", "Bottom bar; labels always visible", "Bottom bar unless width threshold crossed", "Left navigation rail"],
        ["Page padding", "16 px", "16-24 px", "24-32 px with max content width"],
        ["Lists", "Single column", "Single column; denser metadata", "Optional master/detail or two-column cards"],
        ["Forms", "One column; full-width controls", "One column with constrained width", "Two columns only for naturally paired short fields"],
        ["Dashboard", "Stack cards", "Two metric columns", "Two/three columns with stable reading order"],
        ["Map", "Full body; bottom sheet", "Full body; compact sheet", "Map plus side detail panel where helpful"],
        ["Dialogs", "Prefer full-screen/bottom sheet for complex tasks", "Bottom sheet or dialog", "Centered dialog/side panel"],
    ], [1600, 2600, 2600, 2560], 8.0)
    doc.add_heading("Device frames for the prototype", level=2)
    bullets(doc, [
        "Primary phone: 390 x 844 logical pixels.", "Small phone stress case: 360 x 800.",
        "Landscape stress case: 844 x 390.", "Tablet: 768 x 1024 and 1024 x 768.",
        "Use Auto Layout/constraints so the same components adapt; avoid designing each size as an unrelated screen.",
    ])

    doc.add_heading("9. Offline, loading, empty, and error states", level=1)
    table(doc, ["State", "UI treatment", "Required user control"], [
        ["Initial loading", "Skeleton or centered progress; preserve shell", "None unless unusually long"],
        ["Refresh", "Keep current content; small refresh indicator", "Pull-to-refresh or retry"],
        ["Empty", "Specific explanation and contextual illustration/icon", "Create, change filter, or go to prerequisite"],
        ["Offline cached", "Persistent subtle offline banner; timestamp cached data", "Retry; continue safe local work"],
        ["Submission queued", "Success-like queued confirmation with sync status", "View queue/retry/cancel when safe"],
        ["Upload failed", "Per-photo error badge; keep preview", "Retry or remove"],
        ["Permission denied", "Explain why and what still works", "Open settings or continue without feature"],
        ["Session expired", "Preserve non-sensitive draft when possible", "Sign in again"],
        ["Server validation", "Inline field message using API detail", "Correct and resubmit"],
        ["Conflict/duplicate", "Explain existing record or newer update", "View existing, refresh, or revise"],
    ], [1700, 4700, 2960], 8.2)

    doc.add_heading("10. Prototype data kit", level=1)
    doc.add_paragraph("Use stable sample data in every frame so reviewers can follow changes across the flow.")
    table(doc, ["Entity", "Sample data"], [
        ["Company", "Northstar Construction Ltd.; tracking 07:00-18:00; near distance 300 m"],
        ["Projects", "Riverside Tower - Running 62%; East Depot - Delayed 38%; Lake Villas - Planned 0%"],
        ["People", "Amina Rahman (Owner); Farhan Ahmed (Project Engineer); Nusrat Jahan (Site Engineer); Imran Hossain (Employee)"],
        ["Task", "Install Level 3 conduit; High; due today 16:00; Imran; In Progress"],
        ["Issue", "Water ingress at north stairwell; Quality Problem; Critical; Assigned"],
        ["Tracking", "Imran - Inside assigned, 2 min ago, +/- 12 m; another worker - Offline, 46 min ago"],
        ["Progress", "2026-07-21; Structure 80%; MEP 45%; Finishing 18%; 6 photos"],
        ["Notification", "Critical issue assigned to you - 4 min ago - unread"],
    ], [1800, 7560], 8.4)
    doc.add_heading("Prototype copy rules", level=2)
    bullets(doc, [
        "Use sentence case and direct verbs: Save project, Start tracking, Submit report.",
        "Name the object in destructive confirmations: Archive Riverside Tower?",
        "State queue behavior precisely: Saved on this device. It will upload when you reconnect.",
        "For tracking, state purpose, collection window, visibility, retention, and how to stop or revoke.",
        "Avoid blame language for GPS or connection problems; show last known time and recovery action.",
    ])

    doc.add_heading("11. Figma build procedure", level=1)
    bullets(doc, [
        "Create pages named 00 Cover, 01 Foundations, 02 Components, 03 Auth, 04 Owner, 05 HR, 06 Engineer, 07 Employee, 08 Responsive, 09 Archive.",
        "Create color, typography, spacing, radius, and elevation variables/styles before drawing screens.",
        "Build the app shell, app bar, buttons, fields, list row, card, StatusBadge, banners, photo tile, and map controls as components with variants.",
        "Make a grayscale version of the five critical flows and test navigation before visual polish.",
        "Apply the sample data kit consistently; create component properties for role, status, priority, enabled/loading, and network state.",
        "Connect frames with Tap interactions, overlays for confirmations/bottom sheets, and short Smart Animate transitions only where meaning improves.",
        "Set a named starting point for each role flow and share one prototype link with the review instructions.",
        "Preview on a real phone, then review at small-phone, landscape, and tablet frames.",
        "Record decisions in a short changelog and move discarded explorations to the Archive page.",
    ], numbered=True)
    doc.add_heading("Copy-paste brief for Figma Make or a designer", level=2)
    callout(doc, "Prototype prompt", "Design a Material 3 mobile app called SCFMS for construction field management. Use safety orange #EF6C00 with blue-grey structure, 8/12/16 px radii, outdoor-readable hierarchy, and role-based navigation. Create phone and tablet layouts for login, role dashboard, projects with OpenStreetMap polygons, team, tasks, issues, consent-based tracking, team map, daily progress reports, photo uploads, notifications, profile, and company settings. Include loading, empty, offline, queued, validation, permission-denied, and stale-location states. Use the sample Northstar Construction data from this specification. Do not add payroll, inventory, machinery tracking, chat, Google Maps, or ERP features.")

    doc.add_heading("12. Flutter handoff and acceptance checklist", level=1)
    doc.add_heading("Architecture alignment", level=2)
    bullets(doc, [
        "Keep feature-first folders with data, BLoC, presentation, and dependency-registration layers.",
        "Use the existing shared AppTheme, spacing/breakpoint tokens, responsive scaffold, loading/error views, status badge, and role navigation shell.",
        "Keep top-level GoRouter destinations and use existing in-feature navigation conventions for detail/form pages.",
        "Mirror backend permissions for visibility, but rely on the API for authorization.",
        "Use Dio through ApiClient, secure token storage, get_it injection, and the existing offline/retry services.",
        "Use flutter_map/OpenStreetMap only.",
    ])
    doc.add_heading("Definition of done for each screen", level=2)
    bullets(doc, [
        "Matches the approved component and token system in light and dark themes.",
        "Works at 360 px phone width, representative large phone, landscape, and tablet breakpoints without overflow.",
        "Has loading, empty, API error, offline/stale, and permission-limited states where applicable.",
        "Supports keyboard, safe areas, text scaling, semantic labels, adequate contrast, and minimum touch targets.",
        "Does not expose actions the role lacks and handles backend denial gracefully.",
        "Preserves or explicitly discards form drafts; prevents accidental duplicate submissions.",
        "Includes widget tests for behavior and golden coverage for important responsive states.",
        "Has been checked on a real device for map gestures, camera/photo selection, GPS permissions, foreground tracking, and reconnection behavior when relevant.",
    ])
    doc.add_heading("Open decisions before visual sign-off", level=2)
    table(doc, ["Decision", "Why it matters"], [
        ["Launch market and language", "Privacy copy, date/number formats, legal consent, and localization"],
        ["Android-only or Android+iOS pilot", "Background location, permission education, and QA matrix"],
        ["Brand assets", "Logo, approved orange/blue-grey palette, typography, and imagery"],
        ["Default tracking hours and near distance", "Consent copy and dashboard/map interpretation"],
        ["Progress weighting model", "Report form and dashboard calculations"],
        ["Notification delivery provider", "Permission onboarding, token lifecycle, and deep links"],
        ["Staging/production domains", "Real-device build configuration and end-to-end testing"],
    ], [3300, 6060], 8.5)

    path = OUT / "SCFMS_Mobile_UI_Prototype_Documentation.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    paths = [backend_doc(), mobile_doc()]
    for p in paths:
        print(p)
